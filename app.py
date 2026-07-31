import streamlit as st
from google import genai
import requests
import time
import tempfile
import os
import io
import traceback
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from datetime import datetime
from docx import Document

# Any blocking call to the Gemini API (upload or generate) gets run through this
# so a stalled/hung network call raises a real, catchable error instead of
# freezing the app forever with no feedback to the user.
def run_with_timeout(fn, timeout_seconds, *args, **kwargs):
    with ThreadPoolExecutor(max_workers=1) as ex:
        future = ex.submit(fn, *args, **kwargs)
        try:
            return future.result(timeout=timeout_seconds)
        except FutureTimeoutError:
            raise TimeoutError(
                f"The request to Google's API took longer than {timeout_seconds}s and was "
                "aborted. This is usually a temporary network or API issue — please try again."
            )

st.set_page_config(page_title="Jacobs Practice Generator", page_icon="🩺", layout="wide")

# --- Jacobs School of Medicine design system ------------------------------
# Palette: institutional navy + refined UB blue, warm "exam paper" background,
# muted clinical red (live) / sage (complete) accents.
# Type: Source Serif 4 (headings) + IBM Plex Sans (UI) + IBM Plex Mono (data/citations).
st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Source+Serif+4:opsz,wght@8..60,500;8..60,600;8..60,700&family=IBM+Plex+Sans:wght@400;500;600&family=IBM+Plex+Mono:wght@400;500&display=swap');

        :root {
            --jsm-navy: #0a2240;
            --jsm-blue: #00629b;
            --jsm-paper: #f7f5ef;
            --jsm-rule: #d9d2c2;
            --jsm-pulse: #c1443c;
            --jsm-sage: #3c6e57;
        }

        .stApp {
            background-color: var(--jsm-paper);
            background-image: repeating-linear-gradient(
                to bottom,
                rgba(10, 34, 64, 0.035) 0px,
                rgba(10, 34, 64, 0.035) 1px,
                transparent 1px,
                transparent 29px
            );
        }

        html, body, [class*="css"], .stMarkdown, p, label, span, div {
            font-family: 'IBM Plex Sans', sans-serif;
        }

        h1, h2, h3 {
            font-family: 'Source Serif 4', serif !important;
            color: var(--jsm-navy) !important;
            font-weight: 600 !important;
            letter-spacing: -0.01em;
        }

        .block-container { padding-top: 2rem !important; padding-bottom: 0.5rem !important; }
        h1 { font-size: 1.6rem !important; margin-bottom: 0.1rem !important; padding-bottom: 0rem !important; }
        h3 { font-size: 1.05rem !important; margin-top: 0.1rem !important; margin-bottom: 0.3rem !important; padding-bottom: 0.35rem !important; border-bottom: 1px solid var(--jsm-rule) !important; }
        .stCaption { margin-bottom: 0.2rem !important; font-family: 'IBM Plex Sans', sans-serif !important; }
        hr { margin-top: 0.2rem !important; margin-bottom: 0.2rem !important; border-color: var(--jsm-rule) !important; }
        div[data-testid="stVerticalBlock"] > div { gap: 0.3rem !important; }

        /* Force line/word wrapping in text areas and code blocks */
        code, pre, div[data-baseweb="textarea"] textarea {
            white-space: pre-wrap !important;
            word-wrap: break-word !important;
            overflow-x: hidden !important;
            font-family: 'IBM Plex Mono', monospace !important;
        }

        /* Hero banner + signature ECG pulse-line divider */
        .jsm-hero {
            background: var(--jsm-navy);
            padding: 1.6rem 2rem 0 2rem;
            border-radius: 10px;
            margin-bottom: 1.2rem;
            overflow: hidden;
        }
        .jsm-hero-eyebrow {
            font-family: 'IBM Plex Mono', monospace;
            font-size: 0.7rem;
            letter-spacing: 0.14em;
            color: #7fa8cf;
            text-transform: uppercase;
        }
        .jsm-hero-title {
            font-family: 'Source Serif 4', serif;
            color: #ffffff;
            font-size: 1.85rem;
            margin: 0.35rem 0 0.4rem 0;
            font-weight: 600;
        }
        .jsm-hero-sub {
            color: #cfe0f0;
            font-size: 0.92rem;
            margin: 0 0 1.1rem 0;
            max-width: 640px;
        }
        .jsm-pulse-line {
            height: 18px;
            background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='120' height='18' viewBox='0 0 120 18'%3E%3Cpolyline points='0,9 28,9 35,2 42,16 49,1 56,17 63,9 120,9' fill='none' stroke='%234f83ab' stroke-width='1.4'/%3E%3C/svg%3E");
            background-repeat: repeat-x;
            background-size: 120px 18px;
        }

        /* Bordered containers -> paper cards with left accent */
        div[data-testid="stVerticalBlockBorderWrapper"] {
            border: 1px solid var(--jsm-rule) !important;
            border-left: 3px solid var(--jsm-blue) !important;
            border-radius: 6px !important;
            background: #ffffff;
        }
        div[data-testid="stExpander"] {
            border: 1px solid var(--jsm-rule) !important;
            border-left: 3px solid var(--jsm-blue) !important;
            border-radius: 6px !important;
            background: #ffffff;
        }
        div[data-testid="stCodeBlock"], div[data-testid="stTextArea"] textarea {
            border: 1px solid var(--jsm-rule) !important;
            border-radius: 6px !important;
        }

        /* Buttons */
        button[kind="primary"] {
            background-color: var(--jsm-blue) !important;
            border-color: var(--jsm-blue) !important;
            color: #ffffff !important;
            border-radius: 5px !important;
            font-weight: 600 !important;
            letter-spacing: 0.02em;
        }
        button[kind="primary"]:hover {
            background-color: var(--jsm-navy) !important;
            border-color: var(--jsm-navy) !important;
        }
        button[kind="secondary"] {
            border-radius: 5px !important;
            border-color: var(--jsm-navy) !important;
            color: var(--jsm-navy) !important;
        }

        /* Progress bar in institutional blue */
        div[data-testid="stProgress"] > div > div > div {
            background-color: var(--jsm-blue) !important;
        }

        div[data-testid="stAlert"] {
            border-radius: 6px !important;
        }

        /* Status pill used for the "live quiz" badge */
        .jsm-pill {
            display: inline-block;
            font-family: 'IBM Plex Mono', monospace;
            font-size: 0.78rem;
            padding: 3px 12px;
            border-radius: 20px;
            background: #eaf2ee;
            color: var(--jsm-sage);
            border: 1px solid #cfe3d8;
        }

        .jsm-footer {
            margin-top: 1.5rem;
            padding-top: 0.6rem;
            border-top: 1px solid var(--jsm-rule);
            font-family: 'IBM Plex Mono', monospace;
            font-size: 0.72rem;
            color: #8a8578;
        }
    </style>
""", unsafe_allow_html=True)

GEMINI_KEY = str(st.secrets.get("GEMINI_API_KEY", "")).strip()
BIN_ID = str(st.secrets.get("JSONBIN_BIN_ID", "")).strip()
JSONBIN_KEY = str(st.secrets.get("JSONBIN_API_KEY", "")).strip()

if not GEMINI_KEY or GEMINI_KEY == "None":
    st.error("🔑 `GEMINI_API_KEY` is missing in your Streamlit Secrets dashboard!")
    st.stop()

if not BIN_ID or not JSONBIN_KEY:
    st.error("🔑 JSONBin credentials missing in Streamlit Secrets!")
    st.stop()

client = genai.Client(api_key=GEMINI_KEY)

MODEL_CHAIN = [
    "gemini-flash-latest",
    "gemini-2.5-flash",
    "gemini-pro-latest",
    "gemini-1.5-flash"
]

def load_cloud_data():
    url = f"https://api.jsonbin.io/v3/b/{BIN_ID}/latest"
    headers = {"X-Master-Key": JSONBIN_KEY}
    try:
        req = requests.get(url, headers=headers)
        if req.status_code == 200:
            raw_data = req.json().get("record", {})
            return {k: v for k, v in raw_data.items() if isinstance(v, dict) and "sessions" in v}
    except Exception as e:
        st.error(f"Error reading database: {e}")
    return {}

def create_docx(text_content):
    doc = Document()
    doc.add_heading("Jacobs School of Medicine - Practice Questions", level=1)
    for paragraph in text_content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
        else:
            doc.add_paragraph("")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

@st.dialog("⚠️ Clear Generated Quiz?")
def confirm_reset_dialog():
    st.write("Are you sure you want to clear this generated practice quiz?")
    st.warning("⚠️ **Note:** Once cleared, this quiz cannot be restored and you will need to generate a new set.")
    
    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("🚨 Yes, Clear Quiz", type="primary", use_container_width=True):
            st.session_state.generated_quiz = ""
            st.session_state.is_generating = False
            st.rerun()
    with c2:
        if st.button("Cancel", use_container_width=True):
            st.rerun()

st.markdown("""
<div class="jsm-hero">
    <span class="jsm-hero-eyebrow">Jacobs School of Medicine &amp; Biomedical Sciences</span>
    <h1 class="jsm-hero-title">Student Practice Question Generator</h1>
    <p class="jsm-hero-sub">Select your course and lecture sessions to generate practice questions modeled after in-house exam style.</p>
    <div class="jsm-pulse-line"></div>
</div>
""", unsafe_allow_html=True)

data = load_cloud_data()

if not data:
    st.info("No active courses are currently available. Please check back after faculty publish a course.")
    st.stop()

if "is_generating" not in st.session_state:
    st.session_state.is_generating = False
if "generated_quiz" not in st.session_state:
    st.session_state.generated_quiz = ""

has_quiz = bool(st.session_state.generated_quiz)
is_busy = st.session_state.is_generating or has_quiz

st.markdown("---")
col1, col2 = st.columns([1, 1.3], gap="large")

with col1:
    selected_course = st.selectbox(
        "Select Course:", 
        options=list(data.keys()),
        disabled=is_busy
    )
    course_info = data.get(selected_course, {})
    sessions_dict = course_info.get("sessions", {})
    global_course_style = course_info.get("global_style_profile", "")

    if not sessions_dict:
        st.warning(f"No lecture sessions available for '{selected_course}' yet.")
        st.stop()

    def sort_key(item):
        title, details = item
        raw_date = details.get("date", "2099-12-31")
        return (raw_date, title.lower())

    sorted_sessions = sorted(sessions_dict.items(), key=sort_key)

    st.subheader("1. Select Lecture Sessions")
    
    def toggle_all_sessions():
        select_all_state = st.session_state.get(f"select_all_{selected_course}", False)
        for title, _ in sorted_sessions:
            st.session_state[f"cb_{selected_course}_{title}"] = select_all_state

    st.checkbox(
        "Select All Sessions", 
        key=f"select_all_{selected_course}",
        on_change=toggle_all_sessions,
        disabled=is_busy
    )
    
    selected_session_titles = []
    
    with st.container(height=110, border=True):
        for title, details in sorted_sessions:
            raw_date = details.get("date", "")
            formatted_date = ""
            if raw_date:
                try:
                    dt = datetime.strptime(raw_date, "%Y-%m-%d")
                    formatted_date = dt.strftime("%m/%d/%Y") + " - "
                except ValueError:
                    formatted_date = ""
            
            display_label = f"{formatted_date}{title}"
            cb_key = f"cb_{selected_course}_{title}"
            
            if cb_key not in st.session_state:
                st.session_state[cb_key] = False

            if st.checkbox(
                display_label, 
                key=cb_key,
                disabled=is_busy
            ):
                selected_session_titles.append(title)
            
    st.subheader("2. Quiz Parameters")
    
    p_col1, p_col2 = st.columns([1, 1])
    with p_col1:
        num_questions = st.number_input(
            "Questions (1–20):", 
            min_value=1, 
            max_value=20, 
            value=5, 
            step=1,
            disabled=is_busy
        )
    with p_col2:
        arrange_mode = "By Session"
        if len(selected_session_titles) >= 2:
            arrange_mode = st.radio(
                "Arrangement:",
                options=["By Session", "Shuffle"],
                disabled=is_busy
            )
    
    if not is_busy:
        if st.button("🚀 Generate Practice Quiz", type="primary", use_container_width=True):
            if not selected_session_titles:
                st.error("Please select at least one lecture session.")
            else:
                st.session_state.is_generating = True
                st.session_state.generated_quiz = ""
                st.rerun()
    else:
        if st.button("🛑 Clear & Reset Quiz", use_container_width=True):
            confirm_reset_dialog()

with col2:
    st.subheader("3. Generated Quiz Output")
    
    # Priority 1: If an output is generating right now
    if st.session_state.is_generating:
        st.components.v1.html("""
            <script>
            window.addEventListener('beforeunload', function (e) {
                e.preventDefault();
                e.returnValue = '';
            });
            </script>
        """, height=0)

        status_box = st.empty()
        progress_bar = st.progress(0)
        output_container = st.empty()
        
        status_box.info("⚡ Preparing full slide decks for AI context...")
        progress_bar.progress(15)
        
        uploaded_files = []
        combined_styles = ""
        session_instructions = ""
        
        k = len(selected_session_titles)
        base_quota = num_questions // k
        remainder = num_questions % k
        
        try:
            for idx, title in enumerate(selected_session_titles):
                quota = base_quota + (1 if idx < remainder else 0)
                sess = sessions_dict[title]
                slides_text = sess.get("slides", "")

                status_box.info(f"⚡ Uploading slides for session '{title}' ({idx + 1}/{k})...")

                with tempfile.NamedTemporaryFile(mode="w+", delete=False, suffix=".txt") as temp_file:
                    temp_file.write(f"FULL LECTURE SLIDES FOR SESSION: '{title}'\n\n" + slides_text)
                    temp_path = temp_file.name

                try:
                    g_file = run_with_timeout(client.files.upload, 45, file=temp_path)
                finally:
                    os.remove(temp_path)
                uploaded_files.append(g_file)
                
                session_instructions += f"\n- Session '{title}': Generate exactly {quota} question(s)."
                
                sess_style = sess.get("style_profile")
                if sess_style:
                    combined_styles += f"\n--- Faculty Writing Style Guidelines for {title} ---\n" + sess_style
                elif global_course_style:
                    combined_styles += f"\n--- Course-Wide Faculty Writing Style Guidelines ---\n" + global_course_style

            progress_bar.progress(40)
            status_box.info("🧠 Analyzing 100% of slide contents & matching faculty style...")

            prompt = f"""
            You are a medical school faculty member writing completely original, high-yield in-house exam practice questions for students enrolled in {selected_course}.
            
            --- CRITICAL GROUNDING & ORIGINALITY RULES ---
            1. STRICT SCOPE & ORIGINALITY: All questions, options, and distractors MUST be grounded STRICTLY in facts explicitly stated in the provided lecture slide documents. Read through all slides completely. Do NOT copy or closely paraphrase specific questions from existing practice sets.
            2. INDISTINGUISHABLE FACULTY STYLE: Match the faculty's exact tone, clinical vignette complexity, stem phrasing, and distractor design so closely that the AI-generated questions are indistinguishable from real faculty exam questions.
            3. OBJECTIVES ALIGNMENT: Locate the "Session/Lecture Learning Objectives" (usually on early slides) for each session document. Ensure every question directly tests a stated session learning objective.
            4. SINGLE SESSION ASSIGNMENT: Each question corresponds to EXACTLY ONE lecture session document.
            5. ABSOLUTELY NO LATEX: NEVER use LaTeX syntax, math delimiters, or dollar signs (e.g., do NOT use $, $$, \\frac, \\text, \\mathrm, \\mu). Write all numerical values, units, and chemical formulas using plain text and standard characters only (e.g., write "mg/dL", "alpha-1", "H2O", "10-15%", "greater than", "less than").
            
            --- ANSWER KEY RANDOMIZATION & CITATION RULES ---
            6. RANDOMIZED CORRECT ANSWER DISTRIBUTION: You MUST vary the correct answer position randomly across options A, B, C, D, and E. Avoid clustering correct answers on B or C. Distribute correct keys unpredictably across the set (e.g., A, D, E, B, C) so option placement cannot be guessed by students.
            7. SPECIFIC SLIDE NUMBER CITATIONS: In Section 2, every question rationale MUST cite the exact slide/page number where the fact was derived from the document (e.g., "Exact Slide Citation: Slide 14 - Gastric Acid Phase Control").

            --- TARGET QUESTION DISTRIBUTION PER SESSION ---
            {session_instructions}

            --- IN-HOUSE FACULTY QUESTION WRITING STYLE GUIDELINES ---
            Emulate the exact tone, vignette structure, stem phrasing, and distractor style outlined below:
            {combined_styles if combined_styles else "Write clear, high-yield in-house medical school exam questions based strictly on the slides."}

            --- ARRANGEMENT & FORMATTING ---
            * Total Questions to generate: {num_questions}.
            * Arrangement Mode requested: '{arrange_mode}'.
              - If 'By Session': Group all questions for Session 1 together, then Session 2, etc.
              - If 'Shuffle': Interleave and shuffle the questions across the selected sessions randomly.
            
            Format your output clearly into two main sections:
            
            SECTION 1: QUESTIONS
            For each question, output ONLY the clean question header and stem. Do NOT include the session title in Section 1.
            Format:
            Question [Number]
            [Vignette / Stem]
            A) ...
            B) ...
            C) ...
            D) ...
            E) ...

            SECTION 2: ANSWER KEY & RATIONALES
            For each question, explicitly state the corresponding session title, correct letter, rationale, and specific slide number here.
            Format:
            Question [Number]
            - Session: [Session Title]
            - Correct Answer: [Letter A-E]
            - Detailed Rationale: Explaining why the correct option is right based on slide facts, and why each distractor is incorrect.
            - Exact Slide Citation: Slide [Number] ([Slide Topic/Heading])

            --- REQUIRED FOOTER ---
            At the very end of your response, output a blank line followed exactly by:
            Generated by Jacobs Practice Question Generator, in accordance with the JSMBS Generative Artificial Intelligence Use Policy for Medical Students in the Medical Curriculum.
            """

            contents_payload = uploaded_files + [prompt]

            response = None
            
            for target_model in MODEL_CHAIN:
                for attempt in range(3):
                    try:
                        response = run_with_timeout(
                            client.models.generate_content_stream,
                            45,
                            model=target_model,
                            contents=contents_payload
                        )
                        break
                    except Exception as model_err:
                        err_str = str(model_err)
                        if "503" in err_str or "UNAVAILABLE" in err_str or "404" in err_str:
                            status_box.warning(f"⏳ Server demand high on `{target_model}`. Retrying in 2s (Attempt {attempt+1}/3)...")
                            time.sleep(2)
                            continue
                        else:
                            raise model_err
                if response:
                    break

            if not response:
                raise Exception("Google API servers are currently experiencing peak load. Please try again in a few moments.")

            progress_bar.progress(60)
            status_box.info("✍️ Live Streaming: Writing questions & rationales below...")
            
            full_text = ""
            chunk_count = 0
            STALL_TIMEOUT = 60  # seconds with no new chunk before we give up

            def _consume_next(iterator):
                return next(iterator)

            response_iter = iter(response)
            while True:
                try:
                    chunk = run_with_timeout(_consume_next, STALL_TIMEOUT, response_iter)
                except StopIteration:
                    break
                if chunk.text:
                    full_text += chunk.text
                    chunk_count += 1
                    current_prog = min(60 + (chunk_count * 2), 98)
                    progress_bar.progress(current_prog)
                    output_container.text_area("Live Stream Output:", value=full_text, height=450)

            if not full_text.strip():
                raise Exception(
                    "The AI returned an empty response. This can happen if the model "
                    "stalled or was cut off — please try generating again."
                )

            for g_f in uploaded_files:
                try:
                    client.files.delete(name=g_f.name)
                except Exception:
                    pass

            st.session_state.generated_quiz = full_text
            st.session_state.is_generating = False
            st.rerun()

        except Exception as e:
            # Always log the full traceback server-side. When something goes wrong
            # here the browser often shows nothing useful (or nothing at all if the
            # connection stalled), so this is frequently the ONLY place the real
            # cause is visible — check "Manage app" > Logs on Streamlit Cloud.
            print(f"[{datetime.now().isoformat()}] Quiz generation failed: {e}")
            traceback.print_exc()

            for g_f in uploaded_files:
                try:
                    client.files.delete(name=g_f.name)
                except Exception:
                    pass

            st.session_state.is_generating = False
            status_box.empty()
            progress_bar.empty()
            st.error(f"Error generating questions: {e}")

    # Priority 2: Display active, completed quiz output
    elif st.session_state.generated_quiz:
        st.markdown('<span class="jsm-pill">● Practice quiz active</span>', unsafe_allow_html=True)
        st.write("")
        
        tb_col1, tb_col2, tb_col3 = st.columns([6, 1.5, 1.5])
        
        with tb_col2:
            st.download_button(
                label="📄 .TXT",
                data=st.session_state.generated_quiz,
                file_name=f"{selected_course.replace(' ', '_')}_Practice_Quiz.txt",
                mime="text/plain",
                use_container_width=True
            )
            
        with tb_col3:
            docx_data = create_docx(st.session_state.generated_quiz)
            st.download_button(
                label="📝 .DOCX",
                data=docx_data,
                file_name=f"{selected_course.replace(' ', '_')}_Practice_Quiz.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        st.code(st.session_state.generated_quiz, language="markdown")

    # Priority 3: Default initial state
    else:
        st.info("Select options on the left and click 'Generate Practice Quiz' to create questions.")

st.markdown(
    '<div class="jsm-footer">JACOBS SCHOOL OF MEDICINE — GENERATED IN ACCORDANCE WITH THE JSMBS GENERATIVE AI USE POLICY</div>',
    unsafe_allow_html=True
)
